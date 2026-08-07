"""Generate the FaceBench five-model Baseline B Word research paper."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
FIG = ROOT / "paper" / "figures" / "five_models_baseline_b"
OUT = ROOT / "paper" / "manuscript" / "FaceBench-LFW-FiveModels-BaselineB.docx"
MD_OUT = ROOT / "paper" / "manuscript" / "FaceBench-LFW-FiveModels-BaselineB.md"

MODELS = [
    {
        "name": "FaceNet",
        "key": "facenet",
        "acc": 0.9758,
        "prec": 0.9886,
        "rec": 0.9628,
        "f1": 0.9756,
        "auc": 0.9933,
        "eer": 0.0275,
        "eer_th": 0.3527,
        "far": 0.0111,
        "frr": 0.0372,
        "tp": 2874,
        "tn": 2940,
        "fp": 33,
        "fn": 111,
        "scored": 5958,
        "skipped": 42,
    },
    {
        "name": "Dlib",
        "key": "dlib",
        "acc": 0.5009,
        "prec": 0.5009,
        "rec": 1.0000,
        "f1": 0.6675,
        "auc": 0.9922,
        "eer": 0.0234,
        "eer_th": 0.8966,
        "far": 1.0000,
        "frr": 0.0000,
        "tp": 2952,
        "tn": 0,
        "fp": 2941,
        "fn": 0,
        "scored": 5893,
        "skipped": 107,
    },
    {
        "name": "Buffalo-L",
        "key": "buffalo_l",
        "acc": 0.9828,
        "prec": 1.0000,
        "rec": 0.9657,
        "f1": 0.9826,
        "auc": 0.9899,
        "eer": 0.0216,
        "eer_th": 0.1243,
        "far": 0.0000,
        "frr": 0.0343,
        "tp": 2873,
        "tn": 2961,
        "fp": 0,
        "fn": 102,
        "scored": 5936,
        "skipped": 64,
    },
    {
        "name": "AdaFace",
        "key": "adaface",
        "acc": 0.8098,
        "prec": 0.9995,
        "rec": 0.6208,
        "f1": 0.7659,
        "auc": 0.9711,
        "eer": 0.0730,
        "eer_th": 0.1923,
        "far": 0.0003,
        "frr": 0.3792,
        "tp": 1853,
        "tn": 2972,
        "fp": 1,
        "fn": 1132,
        "scored": 5958,
        "skipped": 42,
    },
    {
        "name": "MagFace",
        "key": "magface",
        "acc": 0.8864,
        "prec": 0.9432,
        "rec": 0.8228,
        "f1": 0.8789,
        "auc": 0.9515,
        "eer": 0.1079,
        "eer_th": 0.3433,
        "far": 0.0498,
        "frr": 0.1772,
        "tp": 2456,
        "tn": 2825,
        "fp": 148,
        "fn": 529,
        "scored": 5958,
        "skipped": 42,
    },
]


def _set_run_font(run, *, size=11, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run_font(run, size=14 if level == 1 else 12, bold=True)


def _add_para(doc: Document, text: str, *, bold=False, italic=False, size=11, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def _add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    _set_run_font(run, size=10, italic=True)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, size=9, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            _set_run_font(run, size=9)
    doc.add_paragraph()


def _add_image(doc: Document, path: Path, width_in: float = 5.8) -> None:
    if not path.is_file():
        _add_para(doc, f"[Missing figure: {path.name}]", italic=True, center=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))


def build_docx() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "FaceBench: A Comparative Evaluation of Five Face Recognition Models "
        "on LFW under a Shared RetinaFace Preprocessing Pipeline"
    )
    _set_run_font(run, size=16, bold=True)

    _add_para(
        doc,
        "FaceBench Research Manuscript — Baseline B Multi-Model Study",
        italic=True,
        center=True,
        size=11,
    )
    _add_para(
        doc,
        "Experiment ID: 20260806T192706Z_paper_lfw_five_models_baseline_b_f86260f8  |  "
        "Framework: FaceBench 0.1.0  |  Device: NVIDIA GeForce RTX 3050 (cuda:0)",
        center=True,
        size=10,
    )

    # Abstract
    _add_heading(doc, "Abstract", 1)
    _add_para(
        doc,
        "Fair comparison of face recognition models is often undermined by "
        "inconsistent detection/alignment pipelines, accuracy-only reporting, and "
        "incomplete experimental provenance. We present FaceBench, a reproducible "
        "benchmarking framework, and execute a locked Labeled Faces in the Wild "
        "(LFW) View-2 verification study of five pretrained recognizers—FaceNet, "
        "Dlib, InsightFace Buffalo-L, AdaFace, and MagFace—under a shared RetinaFace/"
        "SCRFD detect-and-crop protocol (Baseline B). All models used cosine "
        "similarity at a fixed decision threshold of 0.40. On up to 5958 of 6000 "
        "protocol pairs, Buffalo-L achieved the highest fixed-threshold accuracy "
        "(0.9828) with zero false accepts, while FaceNet obtained the highest AUC "
        "(0.9933). Dlib exhibited strong ranking quality (AUC 0.9922, EER 0.0234) "
        "but collapsed accuracy at 0.40 due to score-scale mismatch, illustrating "
        "why fixed-threshold accuracy alone is insufficient. AdaFace and MagFace "
        "remained competitive on AUC (0.9711 and 0.9515) under shared "
        "bbox-margin crops that differ from their native alignment recipes. We "
        "release metrics, confusion matrices, ROC/comparison figures, manifests, "
        "and weight checksums to support transparent multi-model evaluation.",
    )
    _add_para(
        doc,
        "Keywords: face recognition; benchmarking; reproducibility; LFW; "
        "RetinaFace; FaceNet; Dlib; Buffalo-L; AdaFace; MagFace; verification metrics",
        italic=True,
        size=10,
    )

    # 1 Intro
    _add_heading(doc, "1. Introduction", 1)
    _add_para(
        doc,
        "Deep face recognition models are routinely compared in research and "
        "industry, yet experimental practice frequently mixes vendor-specific "
        "detectors, different crop geometries, and incomplete reporting of FAR/"
        "FRR, EER, and computational context. Such inconsistencies make it "
        "difficult to attribute performance differences to the recognition "
        "backbone rather than preprocessing.",
    )
    _add_para(
        doc,
        "This paper contributes (i) FaceBench, a modular Python framework for "
        "YAML-driven comparative evaluation, and (ii) a complete Baseline B LFW "
        "study of five locked models under shared RetinaFace preprocessing. "
        "Relative to our prior single-model Baseline A draft (Buffalo-L only), "
        "Baseline B freezes detection/crop once and reuses aligned faces across "
        "all recognizers.",
    )

    # 2 Related
    _add_heading(doc, "2. Related Work", 1)
    _add_para(
        doc,
        "Canonical verification benchmarks such as LFW remain widely used for "
        "controlled pairwise evaluation. Model families including FaceNet "
        "(triplet learning), Dlib’s ResNet descriptor, ArcFace-style InsightFace "
        "packs (Buffalo-L), quality-adaptive margins (AdaFace), and "
        "magnitude-aware embeddings (MagFace) represent distinct training "
        "objectives and input conventions. Tooling for fair comparison must "
        "therefore normalize detection/alignment while still respecting each "
        "model’s embedding interface. FaceBench targets this gap by combining "
        "dataset/model factories, shared detection hooks, multi-axis metrics, "
        "and automated reporting with experiment manifests.",
    )

    # 3 Methods
    _add_heading(doc, "3. Methods", 1)
    _add_heading(doc, "3.1 FaceBench overview", 2)
    _add_para(
        doc,
        "FaceBench organizes evaluation as: configuration → dataset integrity → "
        "identity pairs → shared align (optional) → embedding → matching → "
        "metrics → reports/figures. Recognition adapters remain unchanged in "
        "this study; AdaFace/MagFace official backbones are injected via "
        "paper-local architecture constructors without modifying framework "
        "adapter source.",
    )
    _add_heading(doc, "3.2 Baseline B shared preprocessing", 2)
    _add_para(
        doc,
        "Detection uses InsightFace RetinaFace/SCRFD (buffalo_l detection pack) "
        "at det_size 640×640. Faces are cropped with crop_mode=bbox_margin "
        "(margin 0.35) to preserve context for models that re-detect internally "
        "(Buffalo-L, Dlib) while providing a face-centric crop for resize-based "
        "embedders (FaceNet, AdaFace, MagFace). Alignment is computed once per "
        "unique image and cached across models. Failed detections skip the "
        "affected pair rather than aborting the experiment.",
    )
    _add_heading(doc, "3.3 Models and matching", 2)
    _add_para(
        doc,
        "Models: FaceNet (facenet-pytorch InceptionResnetV1, VGGFace2), Dlib "
        "(5-point landmarks + recognition ResNet), Buffalo-L (InsightFace ONNX), "
        "AdaFace IR-50 MS1MV2, MagFace iResNet50 MS1MV2. Matching uses cosine "
        "similarity with a locked operating threshold of 0.40 for all models. "
        "AdaFace inputs are channel-flipped to BGR in [-1,1]; MagFace uses a "
        "paper-local BGR [0,1] adaptation matching official MagFace inference, "
        "without changing FaceBench preprocess utilities.",
    )
    _add_heading(doc, "3.4 Dataset and protocol", 2)
    rows_proto = [
        ["Dataset", "LFW deepfunneled"],
        ["Protocol", "View-2 pairs.txt (10 folds; 3000 same + 3000 different)"],
        ["Seed", "42"],
        ["Device", "cuda:0 (NVIDIA GeForce RTX 3050)"],
        ["Skip policy", "Skip pairs with missed detection / embed failure"],
        ["Stub backends", "Disabled"],
    ]
    _add_table(doc, ["Factor", "Setting"], rows_proto)
    _add_caption(doc, "Table 1. Locked experimental protocol (Baseline B).")

    # 4 Results
    _add_heading(doc, "4. Results", 1)
    _add_heading(doc, "4.1 Recognition metrics at threshold 0.40", 2)
    metric_rows = []
    for m in MODELS:
        metric_rows.append(
            [
                m["name"],
                f"{m['acc']:.4f}",
                f"{m['prec']:.4f}",
                f"{m['rec']:.4f}",
                f"{m['f1']:.4f}",
                f"{m['auc']:.4f}",
                f"{m['eer']:.4f}",
                f"{m['far']:.4f}",
                f"{m['frr']:.4f}",
                str(m["scored"]),
                str(m["skipped"]),
            ]
        )
    _add_table(
        doc,
        [
            "Model",
            "Acc",
            "Prec",
            "Rec",
            "F1",
            "AUC",
            "EER",
            "FAR",
            "FRR",
            "Scored",
            "Skipped",
        ],
        metric_rows,
    )
    _add_caption(
        doc,
        "Table 2. LFW View-2 verification metrics under Baseline B "
        "(cosine similarity, threshold = 0.40).",
    )

    _add_para(
        doc,
        "Buffalo-L leads fixed-threshold accuracy (0.9828) with FAR = 0 and "
        "EER = 0.0216. FaceNet is close in accuracy (0.9758) and achieves the "
        "best AUC (0.9933). MagFace and AdaFace trail in Acc@0.40 primarily due "
        "to higher FRR under the shared crop and locked threshold, despite "
        "respectable ROC quality. Dlib’s Acc@0.40 ≈ 0.50 with FAR = 1.0 is an "
        "operating-point artifact: its EER threshold (≈0.90) differs sharply "
        "from 0.40, while AUC (0.9922) and EER (0.0234) confirm strong pairwise "
        "ranking.",
    )

    _add_heading(doc, "4.2 Comparison charts", 2)
    _add_image(doc, FIG / "comparison" / "accuracy_comparison.png", 5.6)
    _add_caption(doc, "Figure 1. Accuracy comparison across the five models on LFW.")
    _add_image(doc, FIG / "comparison" / "f1_comparison.png", 5.6)
    _add_caption(doc, "Figure 2. F1-score comparison across the five models on LFW.")
    _add_image(doc, FIG / "comparison" / "radar_comparison.png", 5.2)
    _add_caption(
        doc,
        "Figure 3. Radar comparison of accuracy, F1, AUC, and inverted EER.",
    )

    _add_heading(doc, "4.3 Confusion matrices at threshold 0.40", 2)
    _add_para(
        doc,
        "Confusion matrices below use the locked cosine threshold 0.40. Cell "
        "layout follows FaceBench convention [[TN, FP], [FN, TP]].",
    )
    conf_rows = [
        [
            m["name"],
            str(m["tp"]),
            str(m["tn"]),
            str(m["fp"]),
            str(m["fn"]),
        ]
        for m in MODELS
    ]
    _add_table(doc, ["Model", "TP", "TN", "FP", "FN"], conf_rows)
    _add_caption(doc, "Table 3. Confusion counts at cosine threshold 0.40.")

    for m in MODELS:
        _add_image(
            doc,
            FIG / "confusion" / f"{m['key']}_confusion_matrix.png",
            4.4,
        )
        _add_caption(
            doc,
            f"Figure. Confusion matrix for {m['name']} on LFW (threshold = 0.40).",
        )

    _add_heading(doc, "4.4 ROC curves", 2)
    for m in MODELS:
        _add_image(
            doc,
            FIG / "comparison" / f"{m['key']}_roc_curve.png",
            4.6,
        )
        _add_caption(doc, f"Figure. ROC curve for {m['name']} on LFW View-2.")

    # 5 Discussion
    _add_heading(doc, "5. Discussion", 1)
    _add_para(
        doc,
        "Shared preprocessing narrows a major confound in multi-model face "
        "recognition comparisons. Remaining gaps (AdaFace/MagFace Acc@0.40; "
        "Dlib threshold mismatch) are scientifically informative: they show "
        "that native alignment recipes and score calibrations still matter "
        "even when detection is frozen. Reporting AUC/EER alongside "
        "fixed-threshold accuracy prevents misranking models whose similarity "
        "scales differ (notably Dlib).",
    )
    _add_para(
        doc,
        "Limitations: this study is restricted to LFW View-2; cross-pose, age, "
        "low-resolution, and video datasets in the FaceBench roadmap are not "
        "yet included. Embedding-time figures in the aggregated latency chart "
        "are not used for ranking because the paper runner timed cached "
        "embeddings. MagFace required a paper-local input adaptation to match "
        "official BGR [0,1] inference conventions.",
    )

    # 6 Conclusion
    _add_heading(doc, "6. Conclusion", 1)
    _add_para(
        doc,
        "Under a frozen Baseline B RetinaFace pipeline on LFW, Buffalo-L and "
        "FaceNet provide the strongest overall verification behavior at the "
        "shared 0.40 cosine threshold, with FaceNet leading AUC and Buffalo-L "
        "leading accuracy/FAR. Dlib, AdaFace, and MagFace remain valuable "
        "baselines when interpreted with EER/AUC and model-specific operating "
        "points. FaceBench’s manifests, checksums, confusion matrices, and "
        "comparison figures make this multi-model study fully auditable and "
        "extensible to broader public datasets.",
    )

    # References (short)
    _add_heading(doc, "References", 1)
    refs = [
        "G. B. Huang et al., “Labeled Faces in the Wild,” UMass Tech Report, 2007.",
        "F. Schroff et al., “FaceNet,” CVPR, 2015.",
        "D. E. King, “Dlib-ml,” JMLR, 2009.",
        "J. Deng et al., “ArcFace,” CVPR, 2019; InsightFace model zoo (Buffalo-L).",
        "M. Kim et al., “AdaFace,” CVPR, 2022.",
        "Q. Meng et al., “MagFace,” CVPR, 2021.",
        "J. Deng et al., “RetinaFace,” CVPR, 2020.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Number")
        p.clear()
        run = p.add_run(ref)
        _set_run_font(run, size=10)

    _add_heading(doc, "Appendix: Reproducibility", 1)
    _add_para(
        doc,
        "Primary experiment directory: "
        "experiments/20260806T192706Z_paper_lfw_five_models_baseline_b_f86260f8. "
        "Single comparison report: paper/results/lfw_five_models_baseline_b_report.md. "
        "Weight checksums: paper/weights/checksums.sha256. "
        "Runner: paper/reproduction/run_lfw_five_models_baseline_b.py. "
        "Figures copied to paper/figures/five_models_baseline_b/.",
        size=10,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


def build_markdown() -> Path:
    lines = [
        "# FaceBench: A Comparative Evaluation of Five Face Recognition Models on LFW under a Shared RetinaFace Preprocessing Pipeline",
        "",
        "**Manuscript status:** Complete Baseline B multi-model draft  ",
        "**Experiment ID:** `20260806T192706Z_paper_lfw_five_models_baseline_b_f86260f8`  ",
        "**Framework:** FaceBench 0.1.0  ",
        "**Device:** NVIDIA GeForce RTX 3050 (`cuda:0`)",
        "",
        "## Abstract",
        "",
        "Fair comparison of face recognition models is often undermined by inconsistent "
        "detection/alignment pipelines, accuracy-only reporting, and incomplete experimental "
        "provenance. We present FaceBench and a locked LFW View-2 verification study of "
        "FaceNet, Dlib, Buffalo-L, AdaFace, and MagFace under shared RetinaFace/SCRFD "
        "Baseline B preprocessing with cosine similarity at threshold 0.40. Buffalo-L "
        "achieves the highest Acc@0.40 (0.9828, FAR=0); FaceNet achieves the highest AUC "
        "(0.9933). Dlib shows strong AUC/EER but Acc@0.40 collapse due to score-scale "
        "mismatch. Confusion matrices and comparison charts are included.",
        "",
        "**Keywords:** face recognition, benchmarking, LFW, RetinaFace, FaceNet, Dlib, "
        "Buffalo-L, AdaFace, MagFace",
        "",
        "## Results summary (Table 2)",
        "",
        "| Model | Acc | Prec | Rec | F1 | AUC | EER | FAR | FRR | TP | TN | FP | FN |",
        "|-------|-----|------|-----|----|-----|-----|-----|-----|----|----|----|----|",
    ]
    for m in MODELS:
        lines.append(
            f"| {m['name']} | {m['acc']:.4f} | {m['prec']:.4f} | {m['rec']:.4f} | "
            f"{m['f1']:.4f} | {m['auc']:.4f} | {m['eer']:.4f} | {m['far']:.4f} | "
            f"{m['frr']:.4f} | {m['tp']} | {m['tn']} | {m['fp']} | {m['fn']} |"
        )
    lines.extend(
        [
            "",
            "## Figures",
            "",
            "- Comparison: `paper/figures/five_models_baseline_b/comparison/`",
            "- Confusion matrices: `paper/figures/five_models_baseline_b/confusion/`",
            "- Word manuscript: `paper/manuscript/FaceBench-LFW-FiveModels-BaselineB.docx`",
            "",
        ]
    )
    MD_OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return MD_OUT


def main() -> int:
    docx_path = build_docx()
    md_path = build_markdown()
    print(f"Wrote {docx_path}")
    print(f"Wrote {md_path}")
    print(f"Size_bytes={docx_path.stat().st_size}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
